import docx
import PyPDF2
from tika import parser
import os
import cv2
import numpy as np

class AdvancedDocumentParser:
    def __init__(self):
        pass
    
    def parse_word(self, file_path):
        """增强版Word文档解析"""
        try:
            doc = docx.Document(file_path)
            content = {
                'text': [],
                'tables': [],
                'images': [],
                'headings': [],
                'body_text': []
            }
            
            # 提取文本和标题
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if text.strip():
                    # 判断是否为标题
                    style_name = paragraph.style.name
                    if style_name.startswith('Heading'):
                        content['headings'].append({
                            'text': text,
                            'level': int(style_name.split(' ')[1]) if len(style_name.split(' ')) > 1 else 1
                        })
                    else:
                        content['body_text'].append(text)
            
            # 提取表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                content['tables'].append({
                    'data': table_data,
                    'type': 'table'
                })
            
            # 提取图片
            # 这里使用python-docx的图片提取功能
            # 注意：这需要额外的处理
            
            return content
        except Exception as e:
            return {'error': str(e)}
    
    def parse_pdf(self, file_path):
        """增强版PDF文档解析"""
        try:
            content = {
                'text': [],
                'tables': [],
                'headings': [],
                'body_text': []
            }
            
            # 使用PyPDF2提取文本
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    page_text = page.extract_text()
                    content['text'].append(page_text)
                    
                    # 识别标题和正文
                    if page_text:
                        lines = page_text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line:
                                # 简单的标题识别规则
                                # 1. 全部大写
                                # 2. 以数字或序号开头（如"1. "、"2.1 "）
                                # 3. 长度较短（通常标题不会太长）
                                is_heading = False
                                
                                # 规则1：全部大写
                                if line.isupper():
                                    is_heading = True
                                # 规则2：以数字或序号开头
                                elif len(line) < 100 and (line.startswith(tuple('0123456789')) or '．' in line or '.' in line):
                                    # 检查是否是列表项或标题
                                    if any(pattern in line for pattern in ['章', '节', '条', '款', '项', '目']):
                                        is_heading = True
                                    # 检查是否是类似"1.1 "的格式
                                    elif len(line.split(' ')[0]) < 10 and ('.' in line.split(' ')[0] or '．' in line.split(' ')[0]):
                                        is_heading = True
                                # 规则3：长度较短且可能是标题
                                elif len(line) < 50 and (line.endswith('：') or line.endswith(':')):
                                    is_heading = True
                                
                                if is_heading:
                                    content['headings'].append({
                                        'text': line,
                                        'page': page_num + 1
                                    })
                                else:
                                    content['body_text'].append(line)
            
            # 集成Camelot提取表格
            try:
                import camelot
                tables = camelot.read_pdf(file_path, pages='all', flavor='lattice')
                for i, table in enumerate(tables):
                    content['tables'].append({
                        'data': table.df.to_dict('records'),
                        'type': 'table',
                        'page': table.page
                    })
            except ImportError:
                pass  # 如果没有安装Camelot，跳过表格提取
            
            return content
        except Exception as e:
            return {'error': str(e)}
    
    def parse_image(self, file_path):
        """解析图片文档"""
        try:
            # 读取图像
            image = cv2.imread(file_path)
            if image is None:
                return {'error': '无法读取图像'}
            
            # 转换为灰度图
            if len(image.shape) == 3:
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            else:
                gray = image
            
            # 尝试进行OCR
            try:
                from AdvancedOCRProcessor import AdvancedOCRProcessor
                ocr_processor = AdvancedOCRProcessor()
                ocr_result = ocr_processor.ocr_image(file_path)
                
                content = {
                    'text': [],
                    'headings': [],
                    'body_text': [],
                    'image_info': {
                        'width': image.shape[1],
                        'height': image.shape[0],
                        'channels': image.shape[2] if len(image.shape) == 3 else 1
                    }
                }
                
                if 'text' in ocr_result and ocr_result['text']:
                    content['text'].append(ocr_result['text'])
                    
                    # 识别标题和正文
                    lines = ocr_result['text'].split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            # 简单的标题识别规则
                            is_heading = False
                            
                            # 规则1：全部大写
                            if line.isupper():
                                is_heading = True
                            # 规则2：以数字或序号开头
                            elif len(line) < 100 and (line.startswith(tuple('0123456789')) or '．' in line or '.' in line):
                                # 检查是否是列表项或标题
                                if any(pattern in line for pattern in ['章', '节', '条', '款', '项', '目']):
                                    is_heading = True
                                # 检查是否是类似"1.1 "的格式
                                elif len(line.split(' ')[0]) < 10 and ('.' in line.split(' ')[0] or '．' in line.split(' ')[0]):
                                    is_heading = True
                            # 规则3：长度较短且可能是标题
                            elif len(line) < 50 and (line.endswith('：') or line.endswith(':')):
                                is_heading = True
                            
                            if is_heading:
                                content['headings'].append({
                                    'text': line
                                })
                            else:
                                content['body_text'].append(line)
                
                return content
            except ImportError:
                # 如果没有OCR模块，返回基本信息
                return {
                    'text': [],
                    'image_info': {
                        'width': image.shape[1],
                        'height': image.shape[0],
                        'channels': image.shape[2] if len(image.shape) == 3 else 1
                    }
                }
        except Exception as e:
            return {'error': str(e)}
    
    def parse_document(self, file_path):
        """根据文件类型自动选择解析方法"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.docx':
            return self.parse_word(file_path)
        elif ext == '.pdf':
            return self.parse_pdf(file_path)
        elif ext in ['.jpg', '.jpeg', '.png', '.bmp']:
            return self.parse_image(file_path)
        else:
            # 使用Tika处理其他格式
            try:
                parsed = parser.from_file(file_path)
                return {'text': [parsed.get('content', '')]}
            except Exception as e:
                return {'error': str(e)}